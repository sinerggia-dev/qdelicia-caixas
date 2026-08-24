"""
Gera manual.docx a partir do manual.html.

    python scripts/gerar_docx.py

Por que existe: o manual é escrito uma vez, em HTML, e as outras formas (PDF e Word) saem dele.
Editar o Word à mão faz as três versões divergirem — se precisar mudar conteúdo, mude o
manual.html e rode isto de novo.

Depende de python-docx. O diagrama do ciclo entra como imagem (`scripts/diagrama.png`), gerada à
parte com o Chrome a partir do SVG do manual — o passo a passo está no CLAUDE.md. Se o arquivo não
existir, o script avisa e segue sem ele.

Fontes: usa Segoe UI e Consolas, que existem em qualquer Windows. Archivo e Source Sans 3 são as
do site, mas não estão instaladas na máquina — o Word cairia para uma substituta qualquer.
"""
import io
import os
import re
import sys
from html.parser import HTMLParser

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RAIZ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ENTRADA = os.path.join(RAIZ, 'manual.html')
SAIDA = os.path.join(RAIZ, 'manual.docx')
DIAGRAMA = os.path.join(RAIZ, 'scripts', 'diagrama.png')

# Mesma paleta do site, na versão clara — o papel é sempre claro.
VERDE = RGBColor(0x0F, 0x6A, 0x31)
AMARELO = RGBColor(0x8A, 0x61, 0x00)
VERMELHO = RGBColor(0xA8, 0x1F, 0x1F)
TINTA = RGBColor(0x10, 0x1A, 0x22)
TINTA2 = RGBColor(0x3E, 0x4C, 0x58)
TINTA3 = RGBColor(0x63, 0x73, 0x7E)

FONTE = 'Segoe UI'
MONO = 'Consolas'


# ------------------------------------------------------------------ árvore

class No:
    def __init__(self, tag, attrs=None):
        self.tag = tag
        self.attrs = dict(attrs or {})
        self.filhos = []
        self.texto = ''

    @property
    def classes(self):
        return self.attrs.get('class', '').split()

    def tem(self, c):
        return c in self.classes

    def txt(self):
        """Texto corrido, sem marcação."""
        if self.tag == '#text':
            return self.texto
        return ''.join(f.txt() for f in self.filhos)


VAZIAS = {'br', 'img', 'link', 'meta', 'input', 'hr'}


class Arvore(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.raiz = No('#raiz')
        self.pilha = [self.raiz]

    def handle_starttag(self, tag, attrs):
        n = No(tag, attrs)
        self.pilha[-1].filhos.append(n)
        if tag not in VAZIAS:
            self.pilha.append(n)

    def handle_endtag(self, tag):
        for i in range(len(self.pilha) - 1, 0, -1):
            if self.pilha[i].tag == tag:
                del self.pilha[i:]
                return

    def handle_data(self, data):
        n = No('#text')
        n.texto = data
        self.pilha[-1].filhos.append(n)


def limpar(s):
    return re.sub(r'\s+', ' ', s).strip()


# ------------------------------------------------------------------ estilos

def sombrear(elemento, cor_hex):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), cor_hex)
    elemento.append(shd)


def borda_esquerda(paragrafo, cor_hex, largura=18):
    pPr = paragrafo._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(largura))
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), cor_hex)
    pbdr.append(left)
    pPr.append(pbdr)


def preparar(doc):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    for lado in ('left_margin', 'right_margin'):
        setattr(sec, lado, Cm(2.2))
    sec.top_margin, sec.bottom_margin = Cm(2), Cm(2)

    normal = doc.styles['Normal']
    normal.font.name = FONTE
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TINTA
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15

    for nome, tam, cor in (('Heading 1', 17, VERDE), ('Heading 2', 13, TINTA), ('Heading 3', 11, TINTA2)):
        st = doc.styles[nome]
        st.font.name = FONTE
        st.font.size = Pt(tam)
        st.font.bold = True
        st.font.color.rgb = cor
        st.paragraph_format.space_before = Pt(16 if nome == 'Heading 1' else 12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True


# ------------------------------------------------------------------ inline

def escrever_inline(p, no, negrito=False, italico=False, mono=False, cor=None):
    """Percorre o conteúdo de um nó emitindo runs com a formatação certa."""
    for f in no.filhos:
        if f.tag == '#text':
            t = f.texto
            if not t.strip() and not t.startswith(' '):
                continue
            r = p.add_run(re.sub(r'\s+', ' ', t))
            r.bold = negrito
            r.italic = italico
            r.font.name = MONO if mono else FONTE
            if mono:
                r.font.size = Pt(9.5)
            if cor:
                r.font.color.rgb = cor
        elif f.tag in ('strong', 'b'):
            escrever_inline(p, f, True, italico, mono, cor)
        elif f.tag in ('em', 'i'):
            escrever_inline(p, f, negrito, True, mono, cor)
        elif f.tag == 'code':
            escrever_inline(p, f, negrito, italico, True, cor)
        elif f.tag == 'span' and f.tem('chip'):
            c = AMARELO if f.tem('aguarda') else (VERMELHO if f.tem('alerta') else VERDE)
            r = p.add_run(limpar(f.txt()))
            r.bold = True
            r.font.name = MONO
            r.font.size = Pt(9)
            r.font.color.rgb = c
        elif f.tag == 'a':
            r = p.add_run(limpar(f.txt()))
            r.font.name = MONO
            r.font.size = Pt(9.5)
            r.font.color.rgb = VERDE
        else:
            escrever_inline(p, f, negrito, italico, mono, cor)


def paragrafo(doc, no, estilo=None):
    p = doc.add_paragraph(style=estilo)
    escrever_inline(p, no)
    return p


# ------------------------------------------------------------------ blocos

def tabela(doc, no):
    linhas = []
    cabecalho = []
    for tr in [x for x in descendentes(no, 'tr')]:
        celulas = [c for c in tr.filhos if c.tag in ('th', 'td')]
        if not celulas:
            continue
        if celulas[0].tag == 'th' and not cabecalho:
            cabecalho = celulas
        else:
            linhas.append(celulas)
    if not cabecalho and not linhas:
        return

    ncol = len(cabecalho or linhas[0])
    t = doc.add_table(rows=0, cols=ncol)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    if cabecalho:
        cels = t.add_row().cells
        for i, c in enumerate(cabecalho[:ncol]):
            sombrear(cels[i]._tc.get_or_add_tcPr(), 'EFF3F6')
            p = cels[i].paragraphs[0]
            r = p.add_run(limpar(c.txt()).upper())
            r.bold = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = TINTA3
            r.font.name = FONTE

    for linha in linhas:
        cels = t.add_row().cells
        for i, c in enumerate(linha[:ncol]):
            p = cels[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            escrever_inline(p, c)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def descendentes(no, tag):
    for f in no.filhos:
        if f.tag == tag:
            yield f
        yield from descendentes(f, tag)


def destaque(doc, no):
    """Caixa de regra/aviso: parágrafo sombreado com barra na lateral."""
    vermelha = no.tem('perigo')
    cor_barra = 'A81F1F' if vermelha else 'D9A300'
    fundo = 'FDECEB' if vermelha else 'FFF8E6'

    rot = next((f for f in no.filhos if f.tag == 'span' and f.tem('rot')), None)
    if rot:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        borda_esquerda(p, cor_barra)
        sombrear(p._p.get_or_add_pPr(), fundo)
        r = p.add_run(limpar(rot.txt()).upper())
        r.bold = True
        r.font.size = Pt(8)
        r.font.name = FONTE
        r.font.color.rgb = VERMELHO if vermelha else AMARELO

    for f in no.filhos:
        if f.tag != 'p':
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        borda_esquerda(p, cor_barra)
        sombrear(p._p.get_or_add_pPr(), fundo)
        escrever_inline(p, f)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def acesso(doc, no):
    quem = next((f for f in no.filhos if f.tag == 'span' and f.tem('quem')), None)
    link = next((f for f in no.filhos if f.tag == 'a'), None)
    como = next((f for f in no.filhos if f.tag == 'span' and f.tem('como')), None)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    borda_esquerda(p, '0F6A31')
    if quem:
        r = p.add_run(limpar(quem.txt()))
        r.bold = True
        r.font.name = FONTE
    if link:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(1)
        borda_esquerda(p2, '0F6A31')
        r = p2.add_run(limpar(link.txt()))
        r.font.name = MONO
        r.font.size = Pt(9.5)
        r.font.color.rgb = VERDE
    if como:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(6)
        borda_esquerda(p3, '0F6A31')
        r = p3.add_run(limpar(como.txt()))
        r.font.size = Pt(9.5)
        r.font.color.rgb = TINTA2


def formula(doc, no):
    p = doc.add_paragraph()
    sombrear(p._p.get_or_add_pPr(), 'F5F7F9')
    r = p.add_run(limpar(no.txt()))
    r.font.name = MONO
    r.font.size = Pt(10)


def passos(doc, no):
    """ol.passos — sequência numerada em que cada item tem título e explicação."""
    n = 0
    for li in [f for f in no.filhos if f.tag == 'li']:
        n += 1
        titulo = next(descendentes(li, 'h4'), None)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.6)
        r = p.add_run(f'{n:02d}   ')
        r.bold = True
        r.font.color.rgb = VERDE
        r.font.name = MONO
        if titulo:
            r2 = p.add_run(limpar(titulo.txt()))
            r2.bold = True
        for corpo in descendentes(li, 'p'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(1.6)
            p2.paragraph_format.space_after = Pt(6)
            escrever_inline(p2, corpo)


def figura(doc, no):
    if os.path.exists(DIAGRAMA):
        doc.add_picture(DIAGRAMA, width=Cm(16.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        print('  aviso: scripts/diagrama.png nao encontrado, figura omitida')
    cap = next(descendentes(no, 'figcaption'), None)
    if cap:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(limpar(cap.txt()))
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = TINTA2


def lista(doc, no, numerada=False):
    estilo = 'List Number' if numerada else 'List Bullet'
    for li in [f for f in no.filhos if f.tag == 'li']:
        p = doc.add_paragraph(style=estilo)
        p.paragraph_format.space_after = Pt(3)
        escrever_inline(p, li)


def bloco(doc, no):
    t = no.tag
    if t == 'h2':
        doc.add_heading(limpar(no.txt()), level=1)
    elif t == 'h3':
        doc.add_heading(limpar(no.txt()), level=2)
    elif t == 'h4':
        doc.add_heading(limpar(no.txt()), level=3)
    elif t == 'p':
        paragrafo(doc, no)
    elif t == 'ul':
        lista(doc, no)
    elif t == 'ol':
        passos(doc, no) if no.tem('passos') else lista(doc, no, True)
    elif t == 'figure':
        figura(doc, no)
    elif t == 'table':
        tabela(doc, no)
    elif t == 'div':
        if no.tem('regra'):
            destaque(doc, no)
        elif no.tem('formula'):
            formula(doc, no)
        elif no.tem('acesso'):
            acesso(doc, no)
        else:
            for f in no.filhos:
                bloco(doc, f)
    elif t in ('section', 'main', 'header'):
        for f in no.filhos:
            bloco(doc, f)


# ------------------------------------------------------------------ principal

def main():
    html = io.open(ENTRADA, encoding='utf-8').read()
    corpo = html[html.index('<header class="capa">'):html.index('</body>')]
    a = Arvore()
    a.feed(corpo)

    doc = Document()
    preparar(doc)

    capa = next(descendentes(a.raiz, 'header'), None)
    if capa:
        eyebrow = next((n for n in descendentes(capa, 'p') if n.tem('eyebrow')), None)
        h1 = next(descendentes(capa, 'h1'), None)
        lede = next((n for n in descendentes(capa, 'p') if n.tem('lede')), None)
        if eyebrow:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(limpar(eyebrow.txt()).upper())
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = VERDE
        if h1:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(limpar(h1.txt()))
            r.bold = True
            r.font.size = Pt(26)
            r.font.color.rgb = TINTA
            r.font.name = FONTE
        if lede:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(16)
            r = p.add_run(limpar(lede.txt()))
            r.font.size = Pt(11.5)
            r.font.color.rgb = TINTA2

    main_no = next(descendentes(a.raiz, 'main'), None)
    if main_no is None:
        print('erro: <main> nao encontrado no manual.html')
        return 1
    for sec in [f for f in main_no.filhos if f.tag == 'section']:
        for f in sec.filhos:
            bloco(doc, f)

    doc.save(SAIDA)
    print('gerado:', SAIDA, '-', round(os.path.getsize(SAIDA) / 1024), 'KB')
    return 0


if __name__ == '__main__':
    sys.exit(main())
